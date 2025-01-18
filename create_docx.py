import os
from docx import Document

def create_docx_files(base_path, num_files=6):
    """
    在指定的基础路径下创建指定数量的 .docx 文件，命名为 1.docx, 2.docx, ..., num_files.docx

    :param base_path: 指定的基础文件夹路径
    :param num_files: 要创建的 .docx 文件数量（默认 6）
    """
    # 检查基础路径是否存在
    if not os.path.exists(base_path):
        print(f"指定的基础路径不存在: {base_path}")
        return

    for i in range(1, num_files + 1):
        file_name = f"{i}.docx"
        file_path = os.path.join(base_path, file_name)
        
        # 检查文件是否已经存在
        if os.path.exists(file_path):
            print(f"文件已存在: {file_path}, 跳过创建。")
            continue
        
        # 创建一个新的 Word 文档
        doc = Document()
        doc.add_heading(f"文档 {i}", level=1)
        doc.add_paragraph(f"这是第 {i} 个生成的文档。")
        
        # 保存文档
        try:
            doc.save(file_path)
            print(f"已创建文件: {file_path}")
        except Exception as e:
            print(f"创建文件失败: {file_path}. 错误: {e}")

if __name__ == "__main__":
    # 指定基础文件夹路径
    base_directory = r"******"  # 修改为你的目标路径

    # 调用函数创建文档
    create_docx_files(base_directory, num_files=6)
